"""
Service Availability - Final Report Export (Structured v1)

Builds research-style PDF & DOCX with this order for EACH module:
  [Module Title]
     1) Visual Insight  (chart with caption)
     2) Analysis        (narrative)
     3) CIO Recommendation Tables
          a) Cost Reduction
          b) Performance Improvement
          c) Customer Satisfaction

Also includes:
- Cover page
- KPI summary (in DOCX via table)
- Appendices (optional, e.g. monthly uptime/downtime)

Usage in Streamlit (with st.cache_data):
    from utils_service_availability.report_service import generate_service_availability_report
    pdf_bytes, docx_bytes = generate_service_availability_report(
        df,
        client_name="UEMS",
        period="Jan–Sep 2025",
        logo_path="logo.png",
    )

If you absolutely need a normal function:
    from utils_service_availability.report_service import generate_service_availability_report_fn
"""

from __future__ import annotations
import os, io, re, importlib, tempfile, datetime as dt
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import numpy as np
from PIL import Image
import plotly.graph_objects as go

# DOCX colors + OXML for TOC
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Plotly image save
import plotly.io as pio
pio.kaleido.scope.default_format = "png"
pio.kaleido.scope.default_width = 1280
pio.kaleido.scope.default_height = 720

# PDF (ReportLab)
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# DOCX (optional)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

BRAND = RGBColor(0x00, 0x4C, 0x99)

# =========================
#   Data Structures
# =========================

@dataclass
class FigureBlock:
    title: str
    fig_obj: Any
    caption: str = ""
    analysis_paras: List[str] = field(default_factory=list)
    cio_tables: Dict[str, pd.DataFrame] = field(default_factory=dict)

@dataclass
class ModuleSection:
    name: str
    figures: List[FigureBlock] = field(default_factory=list)
    overview: List[str] = field(default_factory=list)
    cio_tables: Dict[str, pd.DataFrame] = field(default_factory=dict)

    @property
    def analyses(self):
        return self.overview + [p for f in self.figures for p in f.analysis_paras]

@dataclass
class ReportModel:
    title: str
    client_name: str = ""
    period: str = ""
    author: str = ""
    logo_path: Optional[str] = None
    kpis: Dict[str, Any] = field(default_factory=dict)
    modules: List['ModuleSection'] = field(default_factory=list)
    cio_master: Optional[pd.DataFrame] = None
    appendices: Dict[str, pd.DataFrame] = field(default_factory=dict)

# =========================
#   Helpers
# =========================

def _save_fig(fig_obj, path: str):
    """Force Plotly exports (Kaleido) with a corporate palette."""
    try:
        corporate_palette = [
            "#004C99", "#007ACC", "#FF9F1C", "#2ECC71", "#E15F99",
            "#F15BB5", "#00BBF9", "#9B5DE5", "#00F5D4", "#FEE440"
        ]
        fig_obj.update_layout(
            template="plotly_white",
            paper_bgcolor="white",
            plot_bgcolor="white",
            font=dict(color="black"),
            colorway=corporate_palette,
        )
        for i, trace in enumerate(fig_obj.data):
            if "marker" in trace and getattr(trace.marker, "color", None) is None:
                trace.marker.color = corporate_palette[i % len(corporate_palette)]
            if "line" in trace and getattr(trace.line, "color", None) is None:
                trace.line.color = corporate_palette[i % len(corporate_palette)]
        pio.write_image(fig_obj, path, format="png", scale=2, width=1200, height=700)
    except Exception as e:
        print(f"[WARN] Kaleido export issue: {e}")
        try:
            fig_obj.write_image(path, format="png", engine="kaleido", scale=2)
        except Exception:
            try:
                fig_obj.write_html(path.replace(".png", ".html"))
            except Exception as err:
                print(f"[FATAL] Fallback export failed: {err}")

def _md_table_to_df(md: str) -> pd.DataFrame:
    """Parse a markdown pipe table into a DataFrame and clean cell text."""
    lines = [ln.rstrip() for ln in str(md).splitlines() if ln.strip().startswith("|")]
    if len(lines) < 2:
        return pd.DataFrame()

    header = [c.strip() for c in lines[0].strip("|").split("|")]
    data = []
    for ln in lines[2:]:
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if len(cells) == len(header):
            data.append(cells)

    df = pd.DataFrame(data, columns=header)

    # Clean cell text: convert <br>, remove pipes and ugly "— —" patterns
    for col in df.columns:
        df[col] = (
            df[col].astype(str)
                  .str.replace("<br>", " — ", regex=False)
                  .str.replace("|", " ", regex=False)
                  .str.replace(r"[—–]\s*[—–]", " ", regex=True)
                  .str.replace(r"\s+", " ", regex=True)
                  .str.strip()
        )
    return df

def _split_bold_runs(text: str):
    """
    Parse **bold** segments. Returns [(segment, is_bold), ...].
    Also removes literal '|' characters and trims extra spaces.
    Unmatched ** are treated as normal text.
    """
    s = (text or "").replace("|", " ")
    parts = s.split("**")
    out = []
    bold = False
    for seg in parts:
        seg = re.sub(r"\s+", " ", seg).strip()
        if seg:
            out.append((seg, bold))
        bold = not bold

    # If we had an odd number of '**' markers, treat last segment as non-bold
    if len(parts) % 2 != 0 and out:
        out[-1] = (out[-1][0], False)

    return out

# =========================
#   KPI logic (Service Availability)
# =========================

def _compute_overview_kpis(df: pd.DataFrame) -> Dict[str, Any]:
    """
    Align with recommendation_service KPIs:
      - Average Uptime (%)
      - Total Downtime (mins)
      - Total Incidents
      - Total Downtime Cost (RM)
    All values rendered safely as strings for the KPI table.
    """
    d = df.copy()

    # Average uptime (%)
    avg_uptime = np.nan
    if "uptime_percentage" in d.columns:
        vals = pd.to_numeric(d["uptime_percentage"], errors="coerce")
        if vals.notna().any():
            avg_uptime = float(vals.mean())

    # Total downtime (mins)
    total_dt = np.nan
    if "downtime_minutes" in d.columns:
        vals = pd.to_numeric(d["downtime_minutes"], errors="coerce")
        if vals.notna().any():
            total_dt = float(vals.sum())

    # Total incidents
    total_inc = np.nan
    if "incident_count" in d.columns:
        vals = pd.to_numeric(d["incident_count"], errors="coerce")
        if vals.notna().any():
            total_inc = float(vals.sum())

    # Total downtime cost (RM)
    total_cost = np.nan
    if "estimated_cost_downtime" in d.columns:
        vals = pd.to_numeric(d["estimated_cost_downtime"], errors="coerce")
        if vals.notna().any():
            total_cost = float(vals.sum())

    def _fmt(v, fmt, fallback="N/A"):
        if v is None or (isinstance(v, float) and np.isnan(v)):
            return fallback
        try:
            return fmt(v)
        except Exception:
            return fallback

    return {
        "Average Uptime (%)": _fmt(avg_uptime, lambda x: f"{x:.2f}%"),
        "Total Downtime (mins)": _fmt(total_dt, lambda x: f"{x:,.0f}"),
        "Total Incidents": _fmt(total_inc, lambda x: f"{x:,.0f}"),
        "Total Downtime Cost (RM)": _fmt(total_cost, lambda x: f"{x:,.0f}"),
    }

def _build_kpi_figure(kpis: Dict[str, Any]) -> "go.Figure":
    """
    Optional KPI figure (not currently embedded in PDF, but kept for symmetry).
    """
    titles = list(kpis.keys())

    def _clean_val(v):
        if isinstance(v, (int, float)) and pd.notna(v):
            return v
        try:
            vv = float(str(v).replace("%", "").replace(",", ""))
            return vv
        except Exception:
            return str(v)

    values = [_clean_val(kpis.get(t, "N/A")) for t in titles]

    n = len(titles)
    rows = 2
    cols = max(2, (n + 1) // 2)

    fig = go.Figure()
    idx = 0
    for r in range(rows):
        for c in range(cols):
            if idx >= n:
                break
            t = titles[idx]
            v = values[idx]
            fig.add_trace(go.Indicator(
                mode="number",
                value=(v if isinstance(v, (int, float)) else None),
                number={"valueformat": ",.2f"},
                title={"text": t},
                domain={"row": r, "column": c}
            ))
            if not isinstance(v, (int, float)):
                fig.add_annotation(
                    text=str(v),
                    xref="paper", yref="paper",
                    x=(c + 0.5) / cols, y=(1 - r / rows) - 0.2 / rows,
                    showarrow=False, font=dict(size=18)
                )
            idx += 1

    fig.update_layout(
        grid={"rows": rows, "columns": cols, "pattern": "independent"},
        title={"text": "Executive KPI Summary", "x": 0.5},
        margin=dict(l=20, r=20, t=60, b=20),
        template="plotly_white",
        paper_bgcolor="white",
        plot_bgcolor="white",
    )
    return fig

# ---------- Text / header cleanup ----------

_EMOJI_RE = re.compile(
    "[\U0001F600-\U0001F64F"
    "\U0001F300-\U0001F5FF"
    "\U0001F680-\U0001F6FF"
    "\U0001F1E0-\U0001F1FF"
    "\U00002700-\U000027BF"
    "\U0001F900-\U0001F9FF"
    "\U00002600-\U000026FF"
    "]+",
    flags=re.UNICODE,
)

def _sanitize_header(text: str) -> str:
    t = str(text or "")
    t = _EMOJI_RE.sub("", t)
    t = t.replace("|", " ").strip()
    return re.sub(r"\s+", " ", t)

# ---------- Tidy analysis paragraphs ----------

def _tidy_analysis_blocks(raw: str) -> List[str]:
    """
    Normalize long run-on analysis into clean blocks:
    - Keeps **bold** markers
    - Splits common sections and list items
    """
    s = str(raw or "")
    s = s.replace("|", " ")
    s = re.sub(r"\s*-\s+(?=[A-Z])", "\n- ", s)
    s = re.sub(r"\s*•\s+", "\n- ", s)
    for key in [
        "What this graph is:",
        "What it shows in your data:",
        "How to read it operationally:",
        "Why this matters:",
    ]:
        s = s.replace(key, f"\n{key}")
    blocks = [b.strip() for b in s.split("\n") if b.strip()]
    return blocks

_styles = getSampleStyleSheet()
_P = _styles["BodyText"]
_P.fontName = "Helvetica"
_P.fontSize = 10
_P.leading = 14

def _to_html_bold(text: str) -> str:
    """
    Convert **bold** to <b>…</b> and escape &, <, >.
    Remove unmatched '**' and literal pipes.
    """
    s = str(text or "").replace("|", " ").strip()

    def _pair_sub(m):
        inner = m.group(1)
        return "\uE000" + inner + "\uE001"

    s = re.sub(r"\*\*(.+?)\*\*", _pair_sub, s)
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    s = s.replace("\uE000", "<b>").replace("\uE001", "</b>")
    s = s.replace("**", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _split_sentences(text: str) -> List[str]:
    s = str(text or "").strip()
    if not s:
        return []
    if s.startswith("- ") or s.endswith(":"):
        return [s]
    parts = re.split(r'(?<=[.!?])\s+', s)
    return [p.strip() for p in parts if p.strip()]

# =========================
#   PDF Utilities
# =========================

PAGE_W, PAGE_H = A4
M_L, M_R, M_T, M_B = 50, 50, 60, 50
CONTENT_W = PAGE_W - M_L - M_R

def _wrap_text(text: str, c: canvas.Canvas, max_width: float, font="Helvetica", size=10) -> List[str]:
    text = re.sub(r"\r", "", str(text))
    words = text.split()
    lines, cur = [], ""
    for w in words:
        test = (cur + " " + w).strip()
        if c.stringWidth(test, font, size) <= max_width:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines

def _para(c: canvas.Canvas, text: str, x: float, y: float, lh: float = 14, fs: int = 10) -> float:
    c.setFont("Helvetica", fs)
    for line in _wrap_text(text, c, CONTENT_W, size=fs):
        c.drawString(x, y, line)
        y -= lh
        if y < M_B + 40:
            _footer(c)
            c.showPage()
            y = PAGE_H - M_T
            c.setFont("Helvetica", fs)
    return y

def _footer(c: canvas.Canvas):
    c.setStrokeColor(colors.lightgrey)
    c.line(M_L, M_B - 10, PAGE_W - M_R, M_B - 10)
    c.setFont("Helvetica", 8)
    c.drawRightString(PAGE_W - M_R, M_B - 25, f"Page {int(c.getPageNumber())}")

def _para_rich(c: canvas.Canvas, text: str, x: float, y: float, lh: float = 14, fs: int = 10) -> float:
    """Bold-aware paragraph renderer for PDF, word-wrapped."""
    max_w = CONTENT_W
    words_runs = []
    for seg, is_bold in _split_bold_runs(text):
        for w in seg.split(" "):
            if w:
                words_runs.append((w, is_bold))

    line = []
    width = 0.0

    def _font_name(b): return "Helvetica-Bold" if b else "Helvetica"

    for word, is_bold in words_runs:
        word_w = c.stringWidth(word + " ", _font_name(is_bold), fs)
        if width + word_w > max_w and line:
            cur_x = x
            for w, b in line:
                c.setFont(_font_name(b), fs)
                ww = c.stringWidth(w + " ", _font_name(b), fs)
                c.drawString(cur_x, y, w + " ")
                cur_x += ww
            y -= lh
            if y < M_B + 40:
                _footer(c)
                c.showPage()
                y = PAGE_H - M_T
            line, width = [], 0.0
        line.append((word, is_bold))
        width += word_w

    if line:
        cur_x = x
        for w, b in line:
            c.setFont(_font_name(b), fs)
            ww = c.stringWidth(w + " ", _font_name(b), fs)
            c.drawString(cur_x, y, w + " ")
            cur_x += ww
        y -= lh
        if y < M_B + 40:
            _footer(c)
            c.showPage()
            y = PAGE_H - M_T
    return y

def _draw_analysis_rich(c: canvas.Canvas, raw: str, y: float) -> float:
    """
    Render analysis text using bold-aware wrapper, aligned with other templates.
    """
    for block in _tidy_analysis_blocks(raw):
        for sent in _split_sentences(block):
            if not sent:
                continue
            if sent.endswith(":"):
                y = _para_rich(c, sent, M_L, y, lh=14, fs=10)
                y -= 2
                continue
            if sent.startswith("- "):
                y = _para_rich(c, "• " + sent[2:], M_L + 10, y, lh=14, fs=10)
                y -= 2
                continue
            y = _para_rich(c, sent, M_L, y, lh=14, fs=10)
            y -= 2
    return y

def _draw_para_html(c: canvas.Canvas, html: str, x: float, y: float) -> float:
    para = Paragraph(html, _P)
    w, h = para.wrap(CONTENT_W, PAGE_H)
    if h > (y - M_B - 20):
        _footer(c)
        c.showPage()
        y = PAGE_H - M_T
    para.drawOn(c, M_L, y - h)
    return y - h - 2

def _df_to_pdf_table(df: pd.DataFrame, zebra=True) -> Table:
    def _pfrag(txt: str):
        def _mk_para(s: str) -> Paragraph:
            p = Paragraph(s, _P)
            _w, _h = p.wrap(CONTENT_W, PAGE_H)
            if getattr(p, 'blPara', None) is None or (_w == 0 and _h == 0):
                p = Paragraph("&nbsp;", _P)
                p.wrap(CONTENT_W, PAGE_H)
            return p

        try:
            return _mk_para(_to_html_bold(str(txt)))
        except Exception:
            clean = re.sub(r"<[^>]+>", "", str(txt))
            return _mk_para(clean)

    if df is None or df.empty:
        df = pd.DataFrame([["No data"]], columns=["Message"])

    header = [_pfrag(c) for c in list(df.columns)]
    body_rows = []
    for _, row in df.astype(str).iterrows():
        body_rows.append([_pfrag(v) for v in row.tolist()])
    data = [header] + body_rows

    n_cols = len(df.columns)
    col_widths = [CONTENT_W / n_cols] * n_cols if n_cols > 0 else [CONTENT_W]

    tbl = Table(data, colWidths=col_widths, repeatRows=1)
    styles = [
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#F2F2F2")),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('GRID', (0,0), (-1,-1), 0.25, colors.lightgrey),
        ('ALIGN', (0,0), (-1,0), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
    ]
    if zebra:
        styles.append(
            ('ROWBACKGROUNDS', (0,1), (-1,-1),
             [colors.white, colors.HexColor("#FBFBFB")])
        )
    tbl.setStyle(TableStyle(styles))
    return tbl

def _draw_table_paginated(c: canvas.Canvas, df: pd.DataFrame, y: float) -> float:
    if df is None or df.empty:
        df = pd.DataFrame([["No data"]], columns=["Message"])

    rows = df.astype(str).values.tolist()
    header = list(df.columns)
    start = 0

    while start < len(rows):
        block = 30
        while block >= 1:
            chunk = pd.DataFrame(rows[start:start+block], columns=header)
            tbl = _df_to_pdf_table(chunk)
            w, h = tbl.wrapOn(c, CONTENT_W, PAGE_H)

            available = y - M_B - 20
            if h <= available:
                tbl.drawOn(c, M_L, y - h)
                y = y - h - 10
                start += block
                break
            else:
                block -= 1

        if block < 1:
            _footer(c)
            c.showPage()
            y = PAGE_H - M_T
            continue

        if start < len(rows):
            _footer(c)
            c.showPage()
            y = PAGE_H - M_T

    return y

# =========================
#   DOCX Utilities
# =========================

def _docx_brand_heading(paragraph, text, level=1, color=BRAND, size_pt=None, italic=False):
    if level == 1:
        paragraph.style = "Heading 1"; size_pt = size_pt or 16
    elif level == 2:
        paragraph.style = "Heading 2"; size_pt = size_pt or 13
    else:
        paragraph.style = "Heading 3"; size_pt = size_pt or 11
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.font.bold = True if level <= 2 else False
    run.font.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)

def _docx_add_cover_page(doc, model, logo_path="logo.png"):
    """
    DOCX cover page aligned to the PDF front page:
      - Top brand bar with white title (left-aligned)
      - Main title centered mid-page
      - Meta lines centered (Client, Period, Analyst, Generated)
      - Logo centered under meta
      - Footer: 'Confidential — Mesiniaga Berhad 2025'
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]

    # --------------------------
    # 1) Top brand bar (blue)
    # --------------------------
    bar_table = doc.add_table(rows=1, cols=1)
    bar_table.autofit = False
    bar_width = section.page_width - section.left_margin - section.right_margin
    bar_table.columns[0].width = bar_width

    bar_cell = bar_table.rows[0].cells[0]

    # Cell shading = Mesiniaga blue (#004C99)
    tcPr = bar_cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "004C99")
    tcPr.append(shd)

    p_bar = bar_cell.paragraphs[0]
    p_bar.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_bar = p_bar.add_run(str(model.title))
    run_bar.font.size = Pt(18)
    run_bar.font.bold = True
    run_bar.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Small spacing after bar (simulating PDF top band gap)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(40)

    # ---------------------------------
    # 2) Centered main title (mid-page)
    # ---------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(str(model.title))
    run_title.font.size = Pt(18)
    run_title.font.bold = True

    # ---------------------------------
    # 3) Centered meta lines
    # ---------------------------------
    meta_lines = [
        f"Client: {model.client_name}",
        f"Period: {model.period}",
        f"Analyst: {model.author}" if model.author else "",
        f"Generated: {dt.datetime.now().strftime('%d %b %Y')}",
    ]
    for line in [m for m in meta_lines if m]:
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_meta = p_meta.add_run(line)
        r_meta.font.size = Pt(12)

    # ---------------------------------
    # 4) Centered logo underneath meta
    # ---------------------------------
    logo_file = None
    # model.logo_path takes precedence, else fallback to explicit argument
    if model.logo_path and os.path.exists(model.logo_path):
        logo_file = model.logo_path
    elif logo_path and os.path.exists(logo_path):
        logo_file = logo_path

    if logo_file:
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_file, width=Inches(2.0))
        except Exception:
            # If logo fails, just skip silently
            pass

    # ---------------------------------
    # 5) Footer text (same as PDF)
    # ---------------------------------
    try:
        footer = section.footer
        if footer.paragraphs:
            p_footer = footer.paragraphs[0]
        else:
            p_footer = footer.add_paragraph()
        p_footer.text = "Confidential — Mesiniaga Berhad 2025"
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception:
        # footer is non-critical; ignore failures
        pass

    # ---------------------------------
    # 6) Page break (next content starts on page 2)
    # ---------------------------------
    doc.add_page_break()

def _docx_add_toc_placeholder(doc):
    paragraph = doc.add_paragraph()
    _docx_brand_heading(paragraph, "Table of Contents", level=1)
    p = doc.add_paragraph()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fldSimple)
    doc.add_page_break()

def _docx_add_bold_runs(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = (text or "").replace("|", " ").split("**")
    bold = False
    for seg in parts:
        seg = re.sub(r"\s+", " ", seg).strip()
        if not seg:
            bold = not bold
            continue
        run = p.add_run(seg + " ")
        run.bold = bold
        bold = not bold

def _docx_cell_write_bold(cell, text: str):
    cell.text = ""
    parts = (text or "").replace("|", " ").split("**")
    bold = False
    for seg in parts:
        seg = re.sub(r"\s+", " ", seg).strip()
        if not seg:
            bold = not bold
            continue
        run = cell.paragraphs[0].add_run(seg + " ")
        run.bold = bold
        bold = not bold

def _docx_add_analysis_tidy(doc, raw: str):
    blocks = _tidy_analysis_blocks(raw)
    for line in blocks:
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
        elif line.endswith(":"):
            hdr = doc.add_paragraph()
            r = hdr.add_run(line)
            r.bold = True
        else:
            _docx_add_bold_runs(doc, line)

# =========================
#   CAPTURE per module
# =========================

class _DummyExpander:
    def __enter__(self): return self
    def __exit__(self, exc_type, exc, tb): return False

def _capture_module(mod, fn_name: str, df: pd.DataFrame, section_title: str) -> ModuleSection:
    """
    Monkeypatch streamlit + render_cio_tables to capture analysis + CIO tables.
    """
    section = ModuleSection(name=section_title)
    current_fig_idx = -1

    def _add_text(txt: str):
        nonlocal current_fig_idx
        if not txt or not str(txt).strip():
            return
        s = str(txt).strip().replace("```", "").replace("####", "").replace("###", "").strip()
        if current_fig_idx >= 0 and current_fig_idx < len(section.figures):
            tidy = re.sub(r"[ \t]+", " ", s)
            tidy = re.sub(r"\n{2,}", "\n", tidy).strip()
            section.figures[current_fig_idx].analysis_paras.append(tidy)
        else:
            section.overview.append(s)

    class STShim:
        def subheader(self, *a, **k): pass
        def header(self, *a, **k): pass
        def markdown(self, txt, **k): _add_text(txt)
        def write(self, txt): _add_text(txt)
        def info(self, *a, **k): pass
        def warning(self, *a, **k): pass
        def radio(self, label, options, **k): return options[0]
        def expander(self, *a, **k): return _DummyExpander()
        def image(self, *a, **k): pass
        def plotly_chart(self, fig, **k):
            nonlocal current_fig_idx
            try:
                title = (fig.layout.title.text or "").strip()
            except Exception:
                title = ""
            section.figures.append(FigureBlock(title=title or section_title, fig_obj=fig))
            current_fig_idx += 1

    def capture_cio(title, cio):
        target_fb = section.figures[current_fig_idx] if (0 <= current_fig_idx < len(section.figures)) else None

        def _ingest(md: str) -> pd.DataFrame:
            return _md_table_to_df(md) if md else pd.DataFrame()

        if isinstance(cio, dict):
            mapping = {
                "cost": "Cost",
                "performance": "Performance",
                "satisfaction": "Satisfaction",
            }
            for k_in, pillar in mapping.items():
                md = cio.get(k_in)
                if not md:
                    continue
                dfp = _ingest(md)
                if target_fb is not None and not dfp.empty:
                    if pillar in target_fb.cio_tables and not target_fb.cio_tables[pillar].empty:
                        target_fb.cio_tables[pillar] = pd.concat(
                            [target_fb.cio_tables[pillar], dfp],
                            ignore_index=True,
                        )
                    else:
                        target_fb.cio_tables[pillar] = dfp

                if pillar in section.cio_tables and not section.cio_tables[pillar].empty:
                    section.cio_tables[pillar] = pd.concat(
                        [section.cio_tables[pillar], dfp],
                        ignore_index=True,
                    )
                else:
                    section.cio_tables[pillar] = dfp

    orig_st = getattr(mod, 'st', None)
    orig_render = getattr(mod, 'render_cio_tables', None)
    try:
        setattr(mod, 'st', STShim())
        setattr(mod, 'render_cio_tables', capture_cio)
        getattr(mod, fn_name)(df.copy())
    finally:
        if orig_st is not None:
            setattr(mod, 'st', orig_st)
        if orig_render is not None:
            setattr(mod, 'render_cio_tables', orig_render)

    section.overview = [p for p in section.overview if p.strip()]
    for fb in section.figures:
        fb.analysis_paras = [p for p in fb.analysis_paras if p.strip()]

    return section

# =========================
#   PUBLIC: Generate Report
# =========================

def report_service(
    df: pd.DataFrame,
    client_name: str = "",
    period: str = "",
    logo_path: Optional[str] = None,
    author: str = "AI Business Insight Strategist",
    title: str = "Service Availability — Executive Insight Report",
) -> Tuple[bytes, Optional[bytes]]:
    """
    Build a full research-style report (PDF + DOCX) for the Service Availability template.
    """
    modules_cfg = [
        ("utils_service_availability.recommendation_service_availability.executive_summary",  "executive_summary",  "Executive Summary"),
        ("utils_service_availability.recommendation_service_availability.service_overview",   "service_overview",   "Service Overview"),
        ("utils_service_availability.recommendation_service_availability.service_availability","service_availability","Service Availability Metrics"),
        ("utils_service_availability.recommendation_service_availability.historical_availability", "historical_availability", "Historical Availability Trends"),
        ("utils_service_availability.recommendation_service_availability.incident_analysis",  "incident_analysis",  "Incident Analysis"),
        ("utils_service_availability.recommendation_service_availability.planned_maintenance","planned_maintenance","Planned Maintenance and Downtime"),
        ("utils_service_availability.recommendation_service_availability.emergency_changes",  "emergency_changes",  "Emergency Changes and Their Impacts"),
        ("utils_service_availability.recommendation_service_availability.service_recovery",   "service_recovery",   "Service Recovery Time"),
        ("utils_service_availability.recommendation_service_availability.business_impact",    "business_impact",    "Business Impact Analysis"),
        ("utils_service_availability.recommendation_service_availability.resource_utilization","resource_utilization","Resource Utilization and Scalability"),
    ]

    modules: List[ModuleSection] = []

    for mod_path, fn_name, section_title in modules_cfg:
        try:
            mod = importlib.import_module(mod_path)
            section = _capture_module(mod, fn_name, df, section_title)
            modules.append(section)
        except ModuleNotFoundError:
            modules.append(ModuleSection(name=section_title))
            continue
        except Exception:
            modules.append(ModuleSection(name=section_title))
            continue

    kpis = _compute_overview_kpis(df)

    appendices: Dict[str, pd.DataFrame] = {}
    if "created_time" in df.columns:
        tmp = df.copy()
        tmp["created_time"] = pd.to_datetime(tmp["created_time"], errors="coerce")
        tmp = tmp.dropna(subset=["created_time"])
        if not tmp.empty:
            monthly = tmp.copy()
            monthly["month"] = monthly["created_time"].dt.to_period("M").astype(str)

            agg = {}
            if "uptime_percentage" in monthly.columns:
                agg["uptime_percentage"] = lambda s: pd.to_numeric(s, errors="coerce").mean()
            if "downtime_minutes" in monthly.columns:
                agg["downtime_minutes"] = lambda s: pd.to_numeric(s, errors="coerce").sum()
            if "incident_count" in monthly.columns:
                agg["incident_count"] = lambda s: pd.to_numeric(s, errors="coerce").sum()
            if "estimated_cost_downtime" in monthly.columns:
                agg["estimated_cost_downtime"] = lambda s: pd.to_numeric(s, errors="coerce").sum()

            if agg:
                monthly_agg = (
                    monthly.groupby("month", as_index=False)
                           .agg(agg)
                )
                appendices["Monthly Availability & Downtime (table)"] = monthly_agg

    model = ReportModel(
        title=title,
        client_name=client_name,
        period=period,
        author=author,
        logo_path=logo_path if (logo_path and os.path.exists(logo_path)) else None,
        kpis=kpis,
        modules=modules,
        cio_master=None,
        appendices=appendices,
    )

    pdf_bytes = build_pdf(model).getvalue()
    docx_bytes = None
    if DOCX_AVAILABLE:
        try:
            docx_bytes = build_docx(model).getvalue()
        except Exception:
            docx_bytes = None

    return pdf_bytes, docx_bytes

# =========================
#   BUILDERS: PDF / DOCX
# =========================

def build_pdf(model: ReportModel) -> io.BytesIO:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setTitle(model.title)

    # Cover
    c.setFillColor(colors.HexColor("#004C99"))
    c.rect(0, PAGE_H - 60, PAGE_W, 60, stroke=0, fill=1)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(M_L, PAGE_H - 40, _sanitize_header(model.title))

    c.setFillColor(colors.black)
    y = PAGE_H / 2 + 80
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(PAGE_W / 2, y, _sanitize_header(model.title))
    y -= 28

    c.setFont("Helvetica", 12)
    meta = [
        f"Client: {model.client_name}",
        f"Period: {model.period}",
        f"Analyst: {model.author}" if model.author else "",
        f"Generated: {dt.datetime.now().strftime('%d %b %Y')}",
    ]
    for line in [t for t in meta if t]:
        c.drawCentredString(PAGE_W / 2, y, _sanitize_header(line))
        y -= 18

    if model.logo_path and os.path.exists(model.logo_path):
        try:
            img = Image.open(model.logo_path)
            w = 140
            h = int(w * img.height / img.width)
            c.drawImage(ImageReader(img), PAGE_W / 2 - w / 2, y - h - 20, w, h, mask='auto')
            y -= h + 24
        except Exception:
            pass

    _footer(c)
    c.showPage()

    # Sections
    section_num = 2
    with tempfile.TemporaryDirectory() as tmp:
        for idx, module in enumerate(model.modules, start=1):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(M_L, PAGE_H - M_T, f"{section_num}. {module.name}")
            y = PAGE_H - M_T - 18

            c.setFont("Helvetica-Bold", 12)
            c.drawString(M_L, y, f"{section_num}.1 Visual Insight")
            y -= 12

            if module.overview:
                for para in module.overview:
                    y = _draw_analysis_rich(c, para, y)

            for i, fb in enumerate(module.figures):
                if y < 320:
                    _footer(c)
                    c.showPage()
                    y = PAGE_H - M_T
                    c.setFont("Helvetica-Bold", 14)
                    c.drawString(M_L, y, f"{section_num}. {module.name}")
                    y -= 18
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(M_L, y, f"{section_num}.1 Visual Insight")
                    y -= 12

                c.setFont("Helvetica-Bold", 12)
                subtopic_title = f"{section_num}.{i+1} {fb.title or module.name}"
                y = _para(c, subtopic_title, M_L, y, lh=14, fs=12)
                y -= 4

                img_path = os.path.join(tmp, f"m{idx}_fig{i}.png")
                try:
                    _save_fig(fb.fig_obj, img_path)
                    img = Image.open(img_path)
                    tw = CONTENT_W
                    th = int(tw * img.height / img.width)
                    if th > 300:
                        th = 300
                        tw = int(th * img.width / img.height)

                    if y - th < (M_B + 40):
                        _footer(c)
                        c.showPage()
                        y = PAGE_H - M_T
                        c.setFont("Helvetica-Bold", 14)
                        c.drawString(M_L, y, f"{section_num}. {module.name}")
                        y -= 18
                        c.setFont("Helvetica-Bold", 12)
                        c.drawString(M_L, y, f"{section_num}.1 Visual Insight")
                        y -= 12
                        c.setFont("Helvetica-Bold", 12)
                        y = _para(c, subtopic_title, M_L, y, lh=14, fs=12)
                        y -= 4

                    c.drawImage(ImageReader(img), M_L, y - th, tw, th, mask='auto')
                    y -= th + 6
                except Exception:
                    y = _para(c, "[Figure render failed]", M_L, y)

                cap = fb.caption or fb.title or f"{module.name} Chart"
                c.setFont("Helvetica-Oblique", 9)
                y = _para(c, f"Figure {section_num}.{i+1} — {cap}", M_L, y, lh=12, fs=9)
                y -= 6

                if fb.analysis_paras:
                    if y < (M_B + 140):
                        _footer(c)
                        c.showPage()
                        y = PAGE_H - M_T
                        c.setFont("Helvetica-Bold", 14)
                        c.drawString(M_L, y, f"{section_num}. {module.name}")
                        y -= 18
                        c.setFont("Helvetica-Bold", 12)
                        c.drawString(M_L, y, f"{section_num}.1 Visual Insight")
                        y -= 12
                        c.setFont("Helvetica-Bold", 12)
                        y = _para(c, subtopic_title, M_L, y, lh=14, fs=12)
                        y -= 4

                    c.setFont("Helvetica-Bold", 11)
                    y = _para(c, "Analysis", M_L, y, lh=14, fs=11)
                    y -= 2

                    for raw in fb.analysis_paras:
                        y = _draw_analysis_rich(c, raw, y)

                if fb.cio_tables:
                    non_empty = [
                        p for p in ("Cost", "Performance", "Satisfaction")
                        if (fb.cio_tables.get(p) is not None and not fb.cio_tables.get(p).empty)
                    ]
                    for pillar in non_empty:
                        dfp = fb.cio_tables[pillar]
                        if dfp is None or dfp.empty:
                            continue

                        _footer(c)
                        c.showPage()
                        y = PAGE_H - M_T

                        c.setFont("Helvetica-Bold", 14)
                        c.drawString(M_L, y, f"{section_num}. {module.name} — CIO Recommendations")
                        y -= 18
                        c.setFont("Helvetica-Bold", 11)
                        y = _para(c, pillar, M_L, y, lh=14, fs=11)

                        y = _draw_table_paginated(c, dfp, y)

            section_num += 1

    if model.appendices:
        for name, df in (model.appendices or {}).items():
            _footer(c)
            c.showPage()
            y = PAGE_H - M_T

            c.setFont("Helvetica-Bold", 14)
            c.drawString(M_L, y, f"Appendix — {name}")
            y -= 18

            y = _draw_table_paginated(c, df, y)

    _footer(c)
    c.save()
    buf.seek(0)
    return buf

def build_docx(model: ReportModel) -> io.BytesIO:
    doc = Document()
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    except Exception:
        pass

    def add_heading(text, level=1, italic=False):
        p = doc.add_paragraph()
        p.style = f"Heading {level}"
        run = p.add_run(text)
        run.font.color.rgb = BRAND
        run.bold = True if level <= 2 else False
        run.italic = italic
        return p

    sect = doc.sections[0]
    sect.top_margin, sect.bottom_margin = Inches(0.75), Inches(0.75)
    sect.left_margin, sect.right_margin = Inches(0.75), Inches(0.75)

    _docx_add_cover_page(doc, model)
    _docx_add_toc_placeholder(doc)

    add_heading("1. Executive KPI Summary", level=1)
    if model.kpis:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "KPI"
        hdr[1].text = "Value"
        for k, v in model.kpis.items():
            row = table.add_row().cells
            row[0].text = k
            row[1].text = str(v)
    else:
        doc.add_paragraph("No KPI data available.")
    doc.add_page_break()

    section_no = 2
    for idx, module in enumerate(model.modules, start=1):
        add_heading(f"{section_no}. {module.name}", level=1)

        for para in (module.overview or []):
            _docx_add_bold_runs(doc, para)

        with tempfile.TemporaryDirectory() as tmp:
            for f_idx, fig_block in enumerate(module.figures):
                add_heading(f"{section_no}.{f_idx+1} {fig_block.title or module.name}", level=2)

                fig_path = os.path.join(tmp, f"mod{idx}_fig{f_idx}.png")
                try:
                    _save_fig(fig_block.fig_obj, fig_path)
                    doc.add_picture(fig_path, width=Inches(6.0))
                except Exception:
                    doc.add_paragraph("[Figure render failed]")
                cap = f"Figure {section_no}.{f_idx+1} — {fig_block.title or module.name}"
                pcap = doc.add_paragraph(cap)
                pcap.runs[0].italic = True
                pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if fig_block.analysis_paras:
                    add_heading("Analysis", level=3)
                    for para in fig_block.analysis_paras:
                        _docx_add_analysis_tidy(doc, para)

                if fig_block.cio_tables:
                    non_empty = [
                        p for p in ("Cost", "Performance", "Satisfaction")
                        if (fig_block.cio_tables.get(p) is not None and not fig_block.cio_tables.get(p).empty)
                    ]
                    if non_empty:
                        add_heading("CIO Recommendations", level=3)
                        for pillar in non_empty:
                            dfp = fig_block.cio_tables[pillar]
                            add_heading(pillar, level=4, italic=True)
                            table = doc.add_table(rows=1, cols=len(dfp.columns))
                            table.style = "Light Grid Accent 1"
                            hdr_cells = table.rows[0].cells
                            for i, c_name in enumerate(dfp.columns):
                                hdr_cells[i].text = str(c_name)
                            for _, row in dfp.iterrows():
                                row_cells = table.add_row().cells
                                for i, c_name in enumerate(dfp.columns):
                                    _docx_cell_write_bold(row_cells[i], str(row[c_name]))

        doc.add_page_break()
        section_no += 1

    if model.appendices:
        for name, df in model.appendices.items():
            add_heading(f"Appendix — {name}", level=1)
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, c in enumerate(df.columns):
                hdr[i].text = c
            for _, r in df.iterrows():
                row_cells = table.add_row().cells
                for i, c in enumerate(df.columns):
                    row_cells[i].text = str(r[c])
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# =========================
#   Public Callables
# =========================

class _ReportGeneratorCallable:
    """
    Stateless callable wrapper so Streamlit can hash/pickle it when passed
    into @st.cache_data pipelines.
    """
    def __call__(self, df: pd.DataFrame, client_name: str = "", period: str = "", logo_path: Optional[str] = None):
        return report_service(df, client_name=client_name, period=period, logo_path=logo_path)

# Preferred symbol for cached UIs
generate_service_availability_report = _ReportGeneratorCallable()

# Optional plain function alias
def generate_service_availability_report_fn(*args, **kwargs):
    return report_service(*args, **kwargs)
